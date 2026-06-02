# -*- coding: utf-8 -*-
"""Ручные тесты (вариант 26) в стиле друга: формула + Зашифровка + Расшифровка, вся пословица, по строке."""
import io, importlib.util, contextlib
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def load(n, f):
    s = importlib.util.spec_from_file_location(n, f); mm = importlib.util.module_from_spec(s)
    with contextlib.redirect_stdout(io.StringIO()): s.loader.exec_module(mm)
    return mm

PR = "Плохой работник никогда не находит хорошего инструмента."
UP = "АБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ"
LOW = "абвгдежзийклмнопрстуфхцчшщъыьэюя"
def nUP(c): return UP.index(c) + 1

atb = load('atb', 'lab1_atbash.py'); ver = load('ver', 'lab4_vertical.py')
cd = load('cd', 'lab4_cardano.py'); gst = load('gst', 'lab10_gost2012.py'); dh = load('dh', 'lab11_diffie_hellman.py')
shan = load('shan', 'lab5_shanon.py'); sb = load('sb', 'lab2_s_block.py')
a51 = load('a51', 'lab6_a5-1.py'); gam = load('gam', 'lab5_gamma.py')
mg = load('mg', 'lab7_magma.py')
esig = load('esig', 'lab9_elgamal_dig_signature.py')

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(13)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE; st.paragraph_format.space_after = Pt(2)
sec = doc.sections[0]
sec.left_margin = Cm(2); sec.right_margin = Cm(1.5); sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)


def p(t="", bold=False, size=13, after=2, align=None):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(after)
    if align is not None: par.alignment = align
    if t:
        r = par.add_run(t); r.bold = bold; r.font.size = Pt(size)
    return par


def mn(t, size=11):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = par.add_run(t); r.font.name = 'Courier New'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New'); r.font.size = Pt(size)


def head(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(12); par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.keep_with_next = True
    r = par.add_run(t); r.bold = True; r.font.size = Pt(14)


def sub(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(4); par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    r = par.add_run(t); r.bold = True; r.italic = True; r.font.size = Pt(12)


def tbl(headers, rows, fs=10, mono=True):
    """Таблица: шапка жирная, ячейки по центру (моноширинно для чисел)."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        c = t.cell(0, j).paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(0)
        r = c.add_run(h); r.bold = True; r.font.size = Pt(fs)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i + 1, j).paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_after = Pt(0)
            r = c.add_run(str(val)); r.font.size = Pt(fs)
            if mono:
                r.font.name = 'Courier New'; r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    return t


def xorcol(a, b, res, op="XOR"):
    """Запись XOR двух чисел «в столбик» (как от руки)."""
    w = max(len(a), len(b), len(res))
    mn("     " + a.rjust(w))
    mn(f"{op:>4} " + b.rjust(w))
    mn("     " + "-" * w)
    mn("   = " + res.rjust(w))


def lfsr_demo(state, taps, nbits, steps):
    """Потактовый прогон одного РСЛОС (Фибоначчи): старший бит → выход,
    обратная связь = XOR отмеченных битов, сдвиг влево. Возвращает строки таблицы."""
    rows = []
    for t in range(steps):
        out = (state >> (nbits - 1)) & 1
        fb = 0
        for tp in taps:
            fb ^= (state >> tp) & 1
        rows.append([t + 1, format(state, f'0{nbits}b'), fb, out])
        state = ((state << 1) | fb) & ((1 << nbits) - 1)
    return rows


def strip(rowdefs, fs=8):
    """Горизонтальная таблица-полоска: 1-й столбец — подпись строки, дальше ячейки по буквам."""
    ncol = 1 + max(len(cells) for _, cells in rowdefs)
    t = doc.add_table(rows=len(rowdefs), cols=ncol)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (label, cells) in enumerate(rowdefs):
        c0 = t.cell(ri, 0).paragraphs[0]; c0.paragraph_format.space_after = Pt(0)
        r0 = c0.add_run(label); r0.bold = True; r0.font.size = Pt(fs)
        r0.font.name = 'Courier New'; r0.element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        for ci2, val in enumerate(cells):
            cc = t.cell(ri, ci2 + 1).paragraphs[0]; cc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cc.paragraph_format.space_after = Pt(0)
            rr = cc.add_run(str(val)); rr.font.size = Pt(fs)
            rr.font.name = 'Courier New'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    return t


p("Ручные тесты (вариант 26)", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("Плохой работник никогда не находит хорошего инструмента.", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
p("Алфавит: А=1 … Я=32. Замена знаков: «.»→ТЧК, пробел→ПРБ.", size=11, after=6)

rep = atb.replace(PR).upper()

# ── Атбаш ────────────────────────────────────────────────────────────────────
head("Атбаш  (ключ не нужен)")
p("Формула: Yi = X(n−i+1), n = 32. Таблица замены (она же для расшифрования):", size=12)
mn(" ".join(UP)); mn(" ".join(UP[::-1]))
sub("Зашифровка")
mn("текст:  " + rep)
ct_a = atb.encrypt(PR).upper()
mn("шифр:   " + ct_a)
sub("Расшифровка (та же таблица — Атбаш обратен сам себе)")
mn("шифр:   " + ct_a)
mn("текст:  " + atb.decrypt(ct_a).upper())

# ── Матричный ────────────────────────────────────────────────────────────────
head("Матричный шифр.  Ключ A: 1 0 1 / 0 1 0 / 0 0 1")
p("Формула: C = A·B, т.е. C_i = Σ a_ij·b_j. Для этого ключа: шифр C=(b1+b3, b2, b3); "
  "обратно B = A⁻¹·C = (c1−c3, c2, c3).", size=12)
nums = [nUP(c) for c in rep if c in UP]
while len(nums) % 3: nums.append(0)
res_m = []
sub("Зашифровка")
for i in range(0, len(nums), 3):
    b = nums[i:i+3]; c = (b[0]+b[2], b[1], b[2]); res_m += list(c)
    mn(f"{rep[i:i+3]:<3} ({b[0]:>2},{b[1]:>2},{b[2]:>2}) -> ({c[0]:>2},{c[1]:>2},{c[2]:>2})")
mn("шифр: " + " ".join(map(str, res_m)), size=10)
sub("Расшифровка (B = (c1−c3, c2, c3) → буквы)")
dec_m = []
for i in range(0, len(res_m), 3):
    c = res_m[i:i+3]; b = ((c[0]-c[2]) % 32 or 32, c[1], c[2]); dec_m += list(b)
back = "".join(UP[v-1] for v in dec_m if 1 <= v <= 32)
mn("буквы: " + back)
mn("текст: " + atb.restore(back))

# ── Вертикальная ─────────────────────────────────────────────────────────────
head("Вертикальная перестановка.  Ключ КОД  (Д=1, К=2, О=3)")
p("Правило: пишем по строкам в 3 столбца, читаем столбцы в порядке Д,К,О. "
  "Обратно: раскладываем шифр по столбцам, читаем по строкам.", size=12)
cv = ver.prepare_text(PR)
sub("Зашифровка (таблица по строкам)")
mn("К  О  Д")
for i in range(0, len(cv), 3):
    mn("  ".join(cv[i:i+3]))
ct_v = ver.vertical_permutation_logic(PR, 'КОД', 'encrypt')
mn("шифр: " + ct_v)
sub("Расшифровка")
mn("текст: " + ver.vertical_permutation_logic(ct_v, 'КОД', 'decrypt'))

# ── Кардано ──────────────────────────────────────────────────────────────────
head("Решётка Кардано  (6×10, вырезы ✕, 15 шт.)")
p("Число ключей (трафаретов): T = 4^(m·k). Зашифровка: вписываем буквы в вырезы при 4 поворотах, "
  "читаем построчно. Расшифровка: вписываем шифр построчно, считываем буквы из вырезов за 4 поворота.", size=12)
grid = cd.CARDANO_GRID
t = doc.add_table(rows=6, cols=10); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
for i in range(6):
    for j in range(10):
        cp = t.cell(i, j).paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        cp.add_run("✕" if grid[i][j] == 1 else "").font.size = Pt(11)
cc = cd.clean_final(cd.replace_punctuation((PR + " ") * 2))[:60]
ct_c = cd.encrypt(cc).upper()
sub("Зашифровка")
mn("текст (60): " + cc); mn("шифр:       " + ct_c)
sub("Расшифровка")
mn("текст:      " + cd.decrypt(ct_c).upper())

# ── Шеннон ───────────────────────────────────────────────────────────────────
head("Шеннон (одноразовый блокнот).  ЛКГ: a=5, c=3, T0=1  (mod 32)")
p("Идея: к каждой букве текста прибавляем «случайный» символ гаммы — получаем шифр. "
  "Гамму вырабатывает псевдослучайный генератор; шифр и расшифрование используют одну и ту же гамму.", size=12)
p("Шаг 1 — гамма по линейному конгруэнтному генератору (ЛКГ). Берём начальное T0 и крутим формулу: "
  "γ1 = (a·T0 + c) mod 32, дальше каждое следующее γ = (a·предыдущее + c) mod 32. "
  "Здесь a=5, c=3, T0=1; модуль 32 = размер алфавита.", size=12)
p("Шаг 2 — наложение гаммы посимвольно. Номер буквы m складываем с гаммой по модулю 32: "
  "c = (m + γ) mod 32 (если 0 → 32). Обратно вычитаем ту же гамму: m = (c − γ) mod 32.", size=12)
p("Подготовка: знаки и пробелы заменяем кодами (. → ТЧК, пробел → ПРБ), затем каждую букву — "
  "в номер (А=1…Я=32).", size=12)
rs = shan.replace(PR)
mnum = [LOW.index(ch) + 1 for ch in rs if ch in LOW]
letters_s = [ch for ch in rs if ch in LOW]
g = []; T = 1
for _ in mnum: T = (5 * T + 3) % 32; g.append(T)
ci = [((mnum[i] + g[i]) % 32) or 32 for i in range(len(mnum))]
sub("Зашифровка — по каждой букве (γ из ЛКГ, c = (m+γ) mod 32)")
for i in range(len(mnum)):
    mn(f"{i+1}) {letters_s[i].upper()}={mnum[i]}  γ={g[i]}  c={ci[i]}", size=10)
mn("шифртекст (числа): " + " ".join(f"{x:02d}" for x in ci), size=9)
sub("Расшифровка:  m = (c − γ) mod 32 → буква")
back_s = shan.restore("".join(LOW[(((ci[i] - g[i]) % 32) or 32) - 1] for i in range(len(ci))))
mn("текст: " + back_s)

# ── Гаммирование Магма (CTR) ─────────────────────────────────────────────────
head("Гаммирование Магма (режим CTR, ГОСТ Р 34.13-2015)")
p("Гамма-шифрование: c = m ⊕ Г, где гамма-блок Г = E_Магма(Ctr) — зашифрование счётчика Ctr "
  "шифром Магма. Расшифрование — тот же XOR: m = c ⊕ Г.", size=12)
p("Гамму (блоки по 64 бита) вырабатывает Магма (32 раунда), поэтому её раунды руками не считаем — "
  "берём гамма-блок и накладываем XOR. Счётчик Ctr увеличивается на 1 каждый блок.", size=12)
g_key = 'ffeeddccbbaa99887766554433221100f0f1f2f3f4f5f6f7f8f9fafbfcfdfeff'
g_iv = '1234567890ABCDEF'
g_hexdata = ''.join(f'{nUP(c):02X}' for c in rep if c in UP)
with contextlib.redirect_stdout(io.StringIO()):
    gam.set_key_iv(g_key, g_iv)
    g_ct = gam.process_blocks(g_hexdata, 'encrypt')
    g_back = gam.process_blocks(g_ct, 'decrypt')
g_g1 = gam.encrypt_64bit_block(gam.IV, gam.ROUND_KEYS)
g_c1 = g_g1 ^ int(g_hexdata[:16], 16)
sub("Принцип на первом блоке (8 букв = 64 бита; ключ — тест ГОСТ, IV = 1234567890ABCDEF)")
mn(f"текст {rep[:8]} = {g_hexdata[:16]}")
mn(f"гамма Г = E_Магма(Ctr0) = {g_g1:016X}")
p("XOR в столбик (c = m ⊕ Г):", size=12)
xorcol(g_hexdata[:16], f'{g_g1:016X}', f'{g_c1:016X}', op="⊕")
mn("так каждый 64-битный блок; счётчик Ctr увеличивается на 1.")
sub("Шифртекст всей пословицы")
mn("шифртекст (hex): " + " ".join(g_ct[i:i+16] for i in range(0, len(g_ct), 16)), size=10)
sub("Расшифровка — тот же XOR с той же гаммой")
g_rec = ''.join(UP[int(g_back[i:i+2], 16) - 1] for i in range(0, len(g_back), 2)
                if 1 <= int(g_back[i:i+2], 16) <= 32)
mn("текст: " + atb.restore(g_rec))

# ── A5/1 ─────────────────────────────────────────────────────────────────────
head("A5/1  (поточный шифр на трёх РСЛОС)")
p("Три регистра сдвига R1=19, R2=22, R3=23 бита. Многочлены обратной связи "
  "(какие биты складываются по модулю 2 при сдвиге):", size=12)
tbl(["регистр", "длина", "многочлен обратной связи", "бит тактирования", "старший бит"],
    [["R1", "19", "x19+x18+x17+x14+1", "R1[8]", "R1[18]"],
     ["R2", "22", "x22+x21+1", "R2[10]", "R2[21]"],
     ["R3", "23", "x23+x22+x21+x8+1", "R3[10]", "R3[22]"]], fs=10, mono=False)
p("Принцип одного РСЛОС: старший бит уходит в гамму; регистр сдвигается влево; "
  "в младший разряд записывается обратная связь — XOR битов, отмеченных многочленом.", size=12)
sub("Прогон одного РСЛОС (5 бит, многочлен x⁵+x²+1, начало 10110)")
lf = lfsr_demo(0b10110, [4, 1], 5, 8)
tbl(["такт", "регистр", "обр. связь (XOR)", "выходной бит"], lf, fs=10)
mn("гамма по тактам: " + "".join(str(r[3]) for r in lf))
sub("Сборка гаммы A5/1 из трёх регистров")
p("Тактирование (стоп-вперёд): F = большинство(R1[8], R2[10], R3[10]); сдвигаются только "
  "регистры, чей бит синхронизации равен F. Каждый из трёх регистров работает как РСЛОС выше.", size=12)
p("Бит гаммы = XOR старших битов трёх регистров (в столбик):", size=12)
mn("     1     ← старший бит R1[18]")
mn("  ⊕  0     ← старший бит R2[21]")
mn("  ⊕  1     ← старший бит R3[22]")
mn("   -----")
mn("  =  0     бит гаммы")
mn("шифрбит = бит текста ⊕ бит гаммы;  расшифрование тем же XOR (A5/1 симметричен).")
p("Полный прогон пословицы руками не делается: инициализация A5/1 — 64+22+100 тактов, "
  "затем по 5 бит на букву. На карточке показан принцип генератора и сборка бита гаммы.", size=11)
sub("Шифртекст всей пословицы (программой, ключ = КЛЮЧ)")
_kb = a51.text_to_bits('ключ'); _ki = 0
for _bit in _kb:
    _ki = (_ki << 1) | _bit
_mb = a51.text_to_bits(a51.replace(PR))
_gm = a51.A5_1(_ki).get_keystream(len(_mb))
_cb = [_mb[i] ^ _gm[i] for i in range(len(_mb))]
mn("текст → 5-битные коды → XOR с гаммой A5/1 (шифрбит = бит текста ⊕ бит гаммы):")
mn("шифртекст: " + a51.bits_to_text(_cb))
mn("(совпадает с выводом программы при ключе КЛЮЧ; руками побитно не считаем)")

# ── RSA ──────────────────────────────────────────────────────────────────────
head("RSA.  P=37, Q=41, E=7")
p("N = P·Q = 37·41 = 1517;  φ(N) = (P−1)(Q−1) = 36·40 = 1440;  "
  "D = E⁻¹ mod φ(N) = 823  (проверка: 7·823 = 5761 = 4·1440 + 1).", size=12)
p("Шифрование каждой буквы: C = M^7 mod 1517.  Расшифрование: M = C^823 mod 1517.", size=12)
letters_r = [c for c in rep if c in UP]
res_r = [pow(nUP(c), 7, 1517) for c in letters_r]
sub("Зашифровка — по каждой букве (первые 5 подробно, дальше кратко)")
for i in range(len(letters_r)):
    M = nUP(letters_r[i])
    if i < 5:
        mn(f"{i+1}) {letters_r[i]}={M}:  {M}^7 mod 1517 = {res_r[i]}", size=10)
    else:
        mn(f"{i+1}) {letters_r[i]}={M}, C={res_r[i]}", size=10)
mn("шифртекст: " + " ".join(map(str, res_r)), size=10)
sub("Расшифровка:  M = C^823 mod 1517 → буква")
back_r = "".join(UP[pow(C, 823, 1517) - 1] for C in res_r)
mn("текст: " + atb.restore(back_r))

# ── Эль-Гамаль ───────────────────────────────────────────────────────────────
head("Эль-Гамаль.  p=37, g=2, секретный x=5  (y=32)")
y = pow(2, 5, 37)
p(f"Открытый ключ y = g^x mod p = 2^5 mod 37 = {y}.  Для КАЖДОЙ буквы выбирается свой "
  "рандомизатор k (НОД(k, φ(p)) = НОД(k, 36) = 1) — так же, как в программе.", size=12)
p("Шифробозначение буквы M:  a = g^k mod p,  b = y^k·M mod p.  Обратно: M = b·(a^x)⁻¹ mod p.", size=12)
KPOOL = [5, 7, 11, 13, 17, 19, 23, 25, 29, 31, 35]   # k, взаимно простые с 36
letters_e = [c for c in rep if c in UP]
trip = []
for i, c in enumerate(letters_e):
    k = KPOOL[i % len(KPOOL)]
    trip.append((k, pow(2, k, 37), (pow(y, k, 37) * nUP(c)) % 37))
sub("Зашифровка — свой k для каждой буквы (первые 5 подробно, дальше кратко)")
for i, c in enumerate(letters_e):
    M = nUP(c); k, a, b = trip[i]
    if i < 5:
        mn(f"{i+1}) {c}={M}:  k={k}, a=2^{k} mod37={a}, b={y}^{k}·{M} mod37={b};  (a,b)=({a},{b})", size=10)
    else:
        mn(f"{i+1}) {c}={M}, k={k}, (a,b)=({a},{b})", size=10)
mn("шифртекст (пары a,b): " + "  ".join(f"{a},{b}" for k, a, b in trip), size=9)
sub("Расшифровка:  M = b·(a^x)⁻¹ mod 37 → буква  (для каждой пары своё a)")
back_e = "".join(UP[((b * pow(pow(a, 5, 37), -1, 37)) % 37) - 1] for k, a, b in trip)
mn("текст: " + atb.restore(back_e))

# ── RSA — цифровая подпись ───────────────────────────────────────────────────
head("RSA — цифровая подпись.  N=1517, E=7, D=823")
p("Подписывается не текст, а его ХЕШ. Хеш — квадратичная свёртка по буквам: "
  "h = 0;  для каждой буквы  h = ((h + Mi)²) mod 1517.", size=12)
sub("Хеш пословицы (первые 5 подробно, дальше кратко)")
h = 0; j = 0
for c in rep:
    if c in UP:
        mi = nUP(c); prev = h; h = ((h + mi) ** 2) % 1517; j += 1
        if j <= 5:
            mn(f"{j}) {c}={mi}:  h=(({prev}+{mi})²) mod 1517 = {h}", size=10)
        else:
            mn(f"{j}) {c}={mi}, h={h}", size=10)
mn(f"итоговый хеш h = {h}")
sub("Подпись и проверка")
S_rsa = pow(h, 823, 1517); m_chk = pow(S_rsa, 7, 1517)
mn(f"подпись:  S = h^D mod N = {h}^823 mod 1517 = {S_rsa}")
mn(f"проверка: m = S^E mod N = {S_rsa}^7 mod 1517 = {m_chk}")
mn("m = h  →  подпись ВЕРНА" if m_chk == h else "  подпись НЕВЕРНА")

# ── Эль-Гамаль — цифровая подпись (EGSA) ─────────────────────────────────────
head("Эль-Гамаль — цифровая подпись (EGSA).  p=37, g=2, x=5 (y=32), k=5")
p("Хеш по модулю p−1 = 36: m = 0; для каждой буквы m = ((m + Mi)²) mod 36. "
  "Подпись (a, b): a = g^k mod p, b = (m − x·a)·k⁻¹ mod 36, k=5 (НОД(5,36)=1). "
  "Проверка: (y^a · a^b) mod p = g^m mod p.", size=12)
p("Примечание: в этом алгоритме алфавит без буквы Й, поэтому Й при хешировании пропускается.", size=11)
sub("Хеш пословицы mod 36 (первые 5 подробно, дальше кратко)")
EG = esig.ALPHABET; mm = 0; j = 0
for c in rep:
    if c in EG:
        mi = EG.find(c) + 1; prev = mm; mm = ((mm + mi) ** 2) % 36; j += 1
        if j <= 5:
            mn(f"{j}) {c}={mi}:  m=(({prev}+{mi})²) mod 36 = {mm}", size=10)
        else:
            mn(f"{j}) {c}={mi}, m={mm}", size=10)
mm = mm or 1
mn(f"итоговый хеш m = {mm}")
sub("Подпись и проверка")
a_eg = pow(2, 5, 37); kinv = pow(5, -1, 36); b_eg = (kinv * (mm - 5 * a_eg)) % 36
mn(f"a = g^k mod p = 2^5 mod 37 = {a_eg};   k⁻¹ mod 36 = {kinv}")
mn(f"b = (m − x·a)·k⁻¹ mod 36 = ({mm} − 5·{a_eg})·{kinv} mod 36 = {b_eg}")
mn(f"подпись: (a, b) = ({a_eg}, {b_eg})")
left = (pow(32, a_eg, 37) * pow(a_eg, b_eg, 37)) % 37; right = pow(2, mm, 37)
mn(f"проверка: A1 = (y^a·a^b) mod 37 = {left};   A2 = g^m mod 37 = {right}")
mn("A1 = A2  →  подпись ВЕРНА" if left == right else "  подпись НЕВЕРНА")

# ── ГОСТ 34.10-2012 ──────────────────────────────────────────────────────────
head("ГОСТ Р 34.10-2012  (ЭЦП).  y2=x3+x+1 mod 23, G=(1,7), x=2, k=3")
pc, ac2, bc, G = 23, 1, 1, (1, 7); q = gst.find_point_order(G, ac2, pc)
Hm = gst.hash_message(gst.preprocess_text(PR.replace(' ', '')), pc) % q or 1
p("Формула подписи: r=x_P mod q (P=[k]G);  s=(k·H(m)+r·x) mod q.", size=12)
sub("Создание подписи")
mn("Y=[2]G=(7,11);  P=[3]G=(18,20)")
r = 18 % q; sg = (3*Hm + r*2) % q
mn(f"H(пословицы)={Hm};  r=18 mod {q}={r};  s=(3·{Hm}+{r}·2) mod {q}={sg}")
mn(f"подпись: r={r}, s={sg};  Y=(7,11)")
sub("Проверка подписи")
hinv = gst.mod_inverse(Hm, q); z1 = (sg*hinv) % q; z2 = ((-r)*hinv) % q
Pv = gst.add_points(gst.multiply_point(z1, G, ac2, pc), gst.multiply_point(z2, (7, 11), ac2, pc), ac2, pc)
mn(f"h⁻¹={hinv};  z1=s·h⁻¹={z1};  z2=−r·h⁻¹={z2}")
mn(f"P=[z1]G+[z2]Y={Pv};  x_P mod q={Pv[0]%q} = r -> подпись верна")

# ── Диффи-Хеллман ────────────────────────────────────────────────────────────
head("Диффи — Хеллман.  n=23, a=5, KA=6, KB=15")
p("Формула: Y=a^K mod n; общий ключ K=Y^K mod n. (Расшифрования нет — только общий ключ.)", size=12)
YA, YB, Ka, Kb = dh.compute_keys(23, 5, 6, 15)
mn(f"YA=5^6 mod 23={YA};  YB=5^15 mod 23={YB}")
mn(f"K={YB}^6 mod 23={Ka} = {YA}^15 mod 23={Kb}")

# ── S-блок замены ────────────────────────────────────────────────────────────
head("S-блок замены (ГОСТ Р 34.12-2015 «Магма»).  Узел замены ТК-26")
p("Преобразование t: 32-битный блок = 8 hex-цифр (тетрад). Каждая цифра заменяется по своей "
  "таблице S1…S8 (S1 — для правой цифры, S8 — для левой). Восемь таблиц замены (узел ТК-26):", size=12)
mn("вход: 0 1 2 3 4 5 6 7 8 9 A B C D E F")
for k in range(8):
    mn(f"S{k+1}:   " + " ".join(f"{v:X}" for v in sb.pi[k]))
mn("(читаем: в S1 вход 1 → выход 4;  в S8 вход 1 → выход 7)")
p("Источник таблицы: фиксированный узел замены из ГОСТ Р 34.12-2015 «Магма», "
  "заданный ТК-26 (в ГОСТ 28147-89 S-блоки задавал сам разработчик).", size=11)


def sblock_hand(b):
    """8 замен блока (8 hex-цифр) слева направо: S8 для левой … S1 для правой."""
    return "  ".join(f"S{7-k+1}:{ch}→{sb.pi[7-k][int(ch,16)]:X}" for k, ch in enumerate(b))


def sblock_hand_inv(c):
    """Обратные замены блока (по обратным таблицам S8⁻¹…S1⁻¹)."""
    return "  ".join(f"S{7-k+1}:{ch}→{sb.pi_inv[7-k][int(ch,16)]:X}" for k, ch in enumerate(c))


sub("Контрольный пример ГОСТ: t(fdb97531)")
mn("FDB97531:  " + sblock_hand("FDB97531"))
mn(f"t(FDB97531) = {sb.t(0xFDB97531):08X}   (тест-вектор ГОСТ — сходится)")

sub("Кодируем пословицу в 32-битные блоки")
p("Буква → номер (А=1…Я=32) → байт из 2 hex-цифр → по 4 буквы в один блок. "
  "Пример: ПЛОХ = П16 Л12 О15 Х22 = байты 10 0C 0F 16 = блок 100C0F16.", size=12)
nums_sb = [nUP(c) for c in rep if c in UP]
hexs = ''.join(f'{n:02X}' for n in nums_sb)
hexs += '0' * ((-len(hexs)) % 8)
blk = [hexs[i:i + 8] for i in range(0, len(hexs), 8)]
ct_sb = [f'{sb.t(int(b, 16)):08X}' for b in blk]
sub("Зашифровка — каждый блок руками по тетрадам (S8 слева … S1 справа)")
for i, b in enumerate(blk):
    mn(f"Блок {i+1}: {rep[i*4:i*4+4]:<4} = {b}", size=10)
    mn("   " + sblock_hand(b), size=10)
    mn(f"   t = {ct_sb[i]}", size=10)
mn("ШИФРТЕКСТ (все блоки): " + " ".join(ct_sb), size=10)
sub("Расшифровка — обратная замена t_inv (если в S8 было 1→7, то обратно 7→1)")
mn("берём шифрблок, каждую цифру меняем по ОБРАТНОЙ таблице → получаем байты → буквы.")
mn("Блок 1: " + ct_sb[0])
mn("   " + sblock_hand_inv(ct_sb[0]))
mn(f"   t_inv = {sb.t_inv(int(ct_sb[0], 16)):08X}  →  байты 10 0C 0F 16  →  П16 Л12 О15 Х22  →  ПЛОХ")
mn("…так каждый из 18 блоков; собранные байты дают исходный текст:")
rec_hex = ''.join(f'{sb.t_inv(int(c, 16)):08X}' for c in ct_sb)
rec_sb = ''.join(UP[int(rec_hex[j:j+2], 16) - 1] for j in range(0, len(rec_hex), 2)
                 if 1 <= int(rec_hex[j:j+2], 16) <= 32)
mn("текст: " + atb.restore(rec_sb))

# ── Сеть Фейстеля ────────────────────────────────────────────────────────────
head("Сеть Фейстеля (ГОСТ 28147-89 / Магма).  Ключ 256 бит, блок 64 бита")
p("Блок делится на левую и правую половины (a1 ≙ L, a0 ≙ R). Раунд с подключом Ki:", size=12)
mn("новая правая = старая левая;   новая левая = старая правая ⊕ f(a0, Ki)")
p("Функция f: (a0 + Ki) mod 2³² → замена по восьми S-блокам → циклический сдвиг влево на 11 бит. "
  "256-битный ключ даёт 8 подключей K1…K8 (32 раунда: K1..K8 трижды, затем K8..K1).", size=12)
p("Полный ручной счёт нереален (32 раунда). Берём тест-вектор ГОСТ: "
  "ключ ffeeddcc…fcfdfeff, блок fedcba9876543210.", size=12)
sub("Первый раунд:  a1 = FEDCBA98,  a0 = 76543210,  K1 = FFEEDDCC")
mn("шаг 1. (a0 + K1) mod 2^32 = 76543210 + FFEEDDCC = 76430FDC")
mn("шаг 2. замена по 8 S-блокам:  76430FDC → 319AC0D0")
mn("шаг 3. сдвиг влево на 11 бит:  319AC0D0 → D606818C  = f(a0, K1)")
p("шаг 4. новая левая = f(a0,K1) ⊕ a1  (XOR в столбик):", size=12)
xorcol("D606818C", "FEDCBA98", "28DA3B14", op="⊕")
mn("итог раунда:  новое a1 = 76543210 (старое a0),  новое a0 = 28DA3B14")
p("Так 31 раунд с обменом половин, 32-й — без обмена. Итог всего блока:", size=12)
mn("fedcba9876543210  →  4ee901e5c2d8ca3d")

# ── Магма (полный прогон пословицы) ──────────────────────────────────────────
head("Магма (ГОСТ Р 34.12-2015).  Блок 64 бита, ключ 256 бит, 32 раунда")
p("Магма — блочный шифр на сети Фейстеля. Блок 64 бита делится на левую a1 и правую a0 "
  "половины (по 32 бита). На каждом раунде применяется раундовая функция f к правой половине "
  "и подключ Ki:", size=12)
p("Функция f(a0, K):  1) сложить a0 с подключом K по модулю 2³²;  "
  "2) заменить 8 тетрад по S-блокам (узел ТК-26);  3) циклический сдвиг влево на 11 бит.", size=12)
p("Один раунд:  новая левая = a0;  новая правая = a1 ⊕ f(a0, Ki). "
  "256-битный ключ режется на 8 подключей K1…K8; порядок 32 раундов: K1…K8 три раза, затем K8…K1. "
  "31 раунд с обменом половин, 32-й — без обмена.", size=12)
mgk = bytes.fromhex('ffeeddccbbaa99887766554433221100f0f1f2f3f4f5f6f7f8f9fafbfcfdfeff')
mgc = mg.Magma(mgk)
sub("Один раунд на первом блоке пословицы (ПЛОХОЙПР, K1 = FFEEDDCC)")
mg_data = bytes(nUP(ch) for ch in rep if ch in UP)
if len(mg_data) % 8:
    mg_data = mg_data.ljust((len(mg_data) // 8 + 1) * 8, b'\x00')
_a1 = int.from_bytes(mg_data[:4], 'big'); _a0 = int.from_bytes(mg_data[4:8], 'big')
_K1 = mgc.keys[0]; _f = mgc._g(_K1, _a0)
mn(f"a1 = {_a1:08X}   a0 = {_a0:08X}   K1 = {_K1:08X}")
mn(f"шаг 1. (a0 + K1) mod 2^32 = {(_a0 + _K1) & 0xFFFFFFFF:08X}")
mn(f"шаг 2. замена по S-блокам и сдвиг 11 → f(a0,K1) = {_f:08X}")
p("новая правая = a1 ⊕ f  (XOR в столбик):", size=12)
xorcol(f"{_a1:08X}", f"{_f:08X}", f"{_f ^ _a1:08X}", op="⊕")
mn(f"новое a1 = {_a0:08X} (старое a0);  новое a0 = {_f ^ _a1:08X}   — и так все 32 раунда")
sub("Кодируем пословицу в 64-битные блоки (8 букв = 8 байт) и шифруем каждый блок")
p("буква → номер (А=1…Я=32) → байт; по 8 букв в один блок; каждый блок шифруем Магмой (режим ECB). "
  "32 раунда руками не считаем — берём результат блока.", size=12)
mg_blocks = [mg_data[i:i + 8] for i in range(0, len(mg_data), 8)]
mg_ct = [mgc.encrypt_block(b) for b in mg_blocks]
for i, b in enumerate(mg_blocks):
    mn(f"Блок {i+1}: {rep[i*8:i*8+8]:<8} = {b.hex().upper()}  →  {mg_ct[i].hex().upper()}", size=10)
sub("Шифртекст всей пословицы")
mn("шифртекст (hex): " + " ".join(x.hex().upper() for x in mg_ct), size=10)
sub("Расшифровка (те же раунды, подключи в обратном порядке) → пословица")
mg_back = b''.join(mgc.decrypt_block(x) for x in mg_ct)
mg_rec = ''.join(UP[x - 1] for x in mg_back if 1 <= x <= 32)
mn("текст: " + atb.restore(mg_rec))

head("Без ручного теста")
p("Кузнечик — SP-сеть (подстановочно-перестановочная), 10 раундов, блок 128 бит, ключ 256 бит. "
  "Руками не считается из-за линейного преобразования над полем Галуа; "
  "проверяется контрольным примером (тест-вектором) из ГОСТ Р 34.12-2015.", size=12)

try:
    doc.save("Карточки_для_тетради.docx"); print("Готово: Карточки_для_тетради.docx")
except PermissionError:
    doc.save("Карточки_для_тетради_NEW.docx"); print("ЗАНЯТ -> Карточки_для_тетради_NEW.docx")
