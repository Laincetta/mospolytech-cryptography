# -*- coding: utf-8 -*-
"""Сборка Word-отчёта по дисциплине (структура — как у эталона Колесниковой)."""
import os, tokenize, io as _io, keyword as _kw
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROVERB = "Плохой работник никогда не находит хорошего инструмента."
PROVERB_1000 = ((PROVERB + " ") * 18)[:1000]

doc = Document()

# ── Базовый стиль ───────────────────────────────────────────────────────────
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(14)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = st.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)

sec = doc.sections[0]
sec.left_margin = Cm(3); sec.right_margin = Cm(1.5)
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)

# Нумерация страниц (внизу по центру, на титуле не показывается)
sec.different_first_page_header_footer = True
_fp = sec.footer.paragraphs[0]
_fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_r = _fp.add_run(); _r.font.name = 'Times New Roman'; _r.font.size = Pt(12)
_b = OxmlElement('w:fldChar'); _b.set(qn('w:fldCharType'), 'begin')
_i = OxmlElement('w:instrText'); _i.set(qn('xml:space'), 'preserve'); _i.text = 'PAGE'
_e = OxmlElement('w:fldChar'); _e.set(qn('w:fldCharType'), 'end')
for _el in (_b, _i, _e):
    _r._r.append(_el)

# Переопределяем встроенные стили заголовков (для живого оглавления): TNR, чёрный
for hname, hsize in (('Heading 1', 15), ('Heading 2', 14)):
    h = doc.styles[hname]
    h.font.name = 'Times New Roman'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    h.font.size = Pt(hsize); h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12); h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True


def add_toc():
    p = doc.add_paragraph(); run = p.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Оглавление обновится в Word: ПКМ → «Обновить поле» (F9).'
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
    for el in (fc, it, fc2, t, fc3):
        run._r.append(el)


def para(text="", align=None, bold=False, italic=False, size=14, after=6, before=0, first_line=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    if first_line is not None:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p


def heading(text, lvl=1):
    p = doc.add_paragraph(text, style='Heading 1' if lvl == 1 else 'Heading 2')
    p.paragraph_format.keep_with_next = True   # не отрывать заголовок от содержимого
    return p


_KW = set(_kw.kwlist)
_COL = {'kw': '0000C0', 'str': '067D17', 'num': '1750EB',
        'com': '808080', 'def': '00627A', 'call': '795E26'}


def _mono(r):
    r.font.name = 'Courier New'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    r.font.size = Pt(8.5)


def _spans(code):
    """Возвращает {номер_строки: [(c0, c1, цвет, курсив), ...]} — подсветка Python."""
    rows = {}
    try:
        toks = list(tokenize.generate_tokens(_io.StringIO(code).readline))
    except Exception:
        return None
    skip = (tokenize.NL, tokenize.NEWLINE, tokenize.INDENT, tokenize.DEDENT)
    for idx, tok in enumerate(toks):
        tt, ts, (sr, sc), (er, ec), _ = tok
        tname = tokenize.tok_name.get(tt, '')
        color, ital = None, False
        if tt == tokenize.COMMENT:
            color, ital = _COL['com'], True
        elif 'STRING' in tname:
            color = _COL['str']
        elif tt == tokenize.NUMBER:
            color = _COL['num']
        elif tt == tokenize.NAME:
            if ts in _KW:
                color = _COL['kw']
            else:
                j = idx - 1
                while j >= 0 and toks[j].type in skip:
                    j -= 1
                prev = toks[j].string if j >= 0 else ''
                k = idx + 1
                while k < len(toks) and toks[k].type in (tokenize.NL, tokenize.NEWLINE):
                    k += 1
                nxt = toks[k].string if k < len(toks) else ''
                if prev in ('def', 'class'):
                    color = _COL['def']
                elif nxt == '(':
                    color = _COL['call']
        if color is None:
            continue
        for r in range(sr, er + 1):
            cs = sc if r == sr else 0
            ce = ec if r == er else 10 ** 9
            rows.setdefault(r, []).append((cs, ce, color, ital))
    return rows


def code_block(path):
    code = open(path, encoding='utf-8').read().rstrip('\n')
    rows = _spans(code)
    for i, ln in enumerate(code.split('\n'), start=1):
        p = doc.add_paragraph()
        pp = p.paragraph_format
        pp.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pp.space_after = Pt(0); pp.space_before = Pt(0)
        spans = sorted(rows.get(i, [])) if rows else []
        if not ln:
            _mono(p.add_run(' ')); continue
        if not spans:
            _mono(p.add_run(ln)); continue
        pos = 0
        for cs, ce, color, ital in spans:
            cs = max(cs, pos); ce = min(ce, len(ln))
            if ce <= cs:
                continue
            if cs > pos:
                _mono(p.add_run(ln[pos:cs]))
            r = p.add_run(ln[cs:ce]); _mono(r)
            r.font.color.rgb = RGBColor.from_string(color)
            if ital:
                r.italic = True
            pos = ce
        if pos < len(ln):
            _mono(p.add_run(ln[pos:]))


def picture(png):
    if not os.path.exists(png):
        para(f"[схема: {png} — не найдена]", italic=True); return
    w, h = Image.open(png).size
    wcm = 15.5
    hcm = wcm * h / w
    if hcm > 19:                       # высота под страницу (с заголовком влезает целиком)
        hcm = 19.0; wcm = hcm * w / h
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(png, width=Cm(wcm))


def placeholder(text):
    p = doc.add_paragraph()
    r = p.add_run("[ " + text + " ]")
    r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80); r.font.size = Pt(12)


C = WD_ALIGN_PARAGRAPH.CENTER

# ── ТИТУЛЬНЫЙ ЛИСТ ──────────────────────────────────────────────────────────
para("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", C, bold=True, size=12, after=4)
para("Федеральное государственное бюджетное образовательное учреждение высшего образования",
     C, size=11, after=4)
para("«МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»", C, bold=True, size=12, after=10)
para("Факультет информационных технологий", C, size=12, after=2)
para("Кафедра «Информационная безопасность»", C, size=12, after=60)
para("ОТЧЁТ", C, bold=True, size=18, after=4)
para("по дисциплине «Программирование криптографических алгоритмов»", C, bold=True, size=14, after=4)
para("Вариант 26", C, size=14, after=80)
para("Выполнил: студент группы 241-372", after=2)
para("Герасимов Егор Андреевич", after=20)
para("Проверила: Бутакова Н. Г.", after=80)
para("Москва, 2026 г.", C, size=13)
doc.add_page_break()

# ── ОГЛАВЛЕНИЕ ───────────────────────────────────────────────────────────────
heading("Оглавление")
add_toc()
doc.add_page_break()

# ── ВВЕДЕНИЕ / АННОТАЦИЯ ─────────────────────────────────────────────────────
heading("Аннотация")
para("В работе реализован набор криптографических алгоритмов согласно методическим "
     "указаниям ЛАБКРИПТ_2022 (вариант 26). Каждый алгоритм оформлен отдельным модулем "
     "с функциями шифрования и расшифрования, объединёнными единым консольным интерфейсом.",
     first_line=1.25)
para("Среда программирования: PyCharm.", first_line=1.25, after=0)
para("Язык программирования: Python 3.13.", first_line=1.25, after=0)
para("Процедуры для запуска программы: запустить файл «main.py» — единый интерфейс выбора "
     "и запуска любого из реализованных алгоритмов; каждый модуль также запускается "
     "самостоятельно.", first_line=1.25)
para("Пословица-тест (вариант 26):", bold=True, after=0)
para("«" + PROVERB + "»", first_line=1.25)
para("Текст для проверки работы (не менее 1000 знаков):", bold=True, after=0)
para(PROVERB_1000, first_line=1.25)
doc.add_page_break()

# ── ИНТЕРФЕЙС ───────────────────────────────────────────────────────────────
heading("Интерфейс программы")
para("Интерфейс консольный. После запуска файла «main.py» пользователю выводится меню "
     "со списком алгоритмов. По введённому номеру запускается выбранный алгоритм; "
     "выход из программы — по пункту «0». Предусмотрены также служебные команды: прогон "
     "пословицы через все алгоритмы, тест на 1000 знаков и вывод таблицы ключей.",
     first_line=1.25)
heading("Блок-схема программы-интерфейса", 2)
picture("_preview_interface.png")
heading("Код программы-интерфейса", 2)
code_block("main.py")
doc.add_page_break()

# ── ОБЩИЕ ФУНКЦИИ ───────────────────────────────────────────────────────────
heading("Общие функции обработки текста")
para("Перед шифрованием знаки препинания и пробелы заменяются буквенными кодами "
     "(функция «замена»: «.» → «ТЧК», «,» → «ЗПТ», пробел → «ПРБ» и т. д.). После "
     "расшифрования выполняется обратная операция (функция «восстановление»: «ТЧК» → «.» "
     "и т. д.). Это позволяет шифровать произвольный текст в рамках 32-буквенного алфавита.",
     first_line=1.25)
heading("Блок-схема функции замены спецсимволов", 2)
picture("_preview_replace.png")
heading("Блок-схема функции восстановления спецсимволов", 2)
picture("_preview_restore.png")
doc.add_page_break()

# ── АЛГОРИТМЫ ────────────────────────────────────────────────────────────────
D = {
 'atbash': "Шифр простой моноалфавитной замены. Каждая буква заменяется на симметричную "
           "ей относительно центра алфавита: i-я буква открытого текста заменяется буквой "
           "с номером n−i+1, где n = 32 — мощность алфавита. Ключ не требуется "
           "(преобразование самообратное).",
 'sblock': "Узел замены шифра «Магма» (ГОСТ Р 34.12-2015, набор ТК-26). 32-битное слово "
           "делится на восемь 4-битовых тетрад, каждая заменяется по своей таблице из "
           "восьми S-блоков (перестановка чисел 0…15). Контрольный тест-вектор "
           "преобразования t: t(fdb97531) = 2a196f34. Сам циклический сдвиг на 11 бит "
           "входит в раундовую функцию шифра (см. Магму и сеть Фейстеля).",
 'hill':   "Блочный шифр на основе матричной алгебры. Буквы переводятся в числа "
           "(А = 1 … Я = 32), n-грамма открытого текста (вектор-блок B из 3 чисел) "
           "умножается на ключевую матрицу A: вектор шифртекста C = A · B. Расшифрование "
           "выполняется обратной матрицей: B = A⁻¹ · C (определитель матрицы должен быть "
           "отличен от нуля).",
 'vertical': "Шифр перестановки. Открытый текст построчно записывается в прямоугольную "
           "таблицу, ширина которой равна длине ключевого слова. Шифртекст получается "
           "считыванием столбцов в порядке, заданном расположением букв ключа в алфавите. "
           "При расшифровании учитываются длинные и короткие столбцы.",
 'cardano': "Поворотная решётка Кардано — частный случай маршрутной перестановки. "
           "Используется трафарет 6×10 с 15 вырезами; при четырёх поворотах вырезы "
           "полностью покрывают все 60 клеток. Буквы вписываются в вырезы при каждом "
           "повороте, пустые места заполняются «пустышками», шифртекст читается построчно "
           "(60 символов на блок).",
 'feistel': "Сеть Фейстеля (ГОСТ 28147-89 / «Магма»). 64-битный блок делится на старшую и "
           "младшую половины. На каждом из 31 раунда: младшая половина складывается с "
           "подключом по модулю 2³², результат заменяется по S-блокам и циклически "
           "сдвигается влево на 11 бит, складывается со старшей половиной по модулю 2, "
           "после чего половины меняются местами. 32-й раунд — без обмена половин. "
           "256-битный ключ даёт восемь подключей K1…K8.",
 'shannon': "Одноразовый блокнот К. Шеннона. Гамма вырабатывается линейным конгруэнтным "
           "генератором: γ = (a·T + c) mod 32. Шифрсимвол c = (m + γ) mod 32 (при нулевом "
           "результате берётся 32). Условия максимального периода: a ≡ 1 (mod 4), c — "
           "нечётное число.",
 'gamma':  "Режим гаммирования шифра «Магма» (счётчик, CTR; ГОСТ Р 34.13-2015). Гамма "
           "вырабатывается зашифрованием значения счётчика в режиме простой замены и "
           "складывается с блоком данных по модулю 2: c = m ⊕ Eₖ(счётчик). Счётчик "
           "инициализируется синхропосылкой и для каждого блока увеличивается на 1 по "
           "модулю 2⁶⁴. Шифрование и расшифрование выполняются одинаково.",
 'a51':    "Поточный шифр на трёх регистрах сдвига с линейной обратной связью (РСЛОС) "
           "длиной 19, 22 и 23 бита. Инициализация: 64 такта загрузки ключа, 22 такта "
           "загрузки номера кадра, 100 тактов холостого прогона. Тактирование — по "
           "мажоритарной функции (режим «старт-стоп»). Бит гаммы — сумма по модулю 2 "
           "старших битов трёх регистров; гамма накладывается на текст по модулю 2.",
 'magma':  "Блочный шифр «Магма» (ГОСТ Р 34.12-2015) на основе сети Фейстеля. Блок 64 бита, "
           "256-битный ключ разворачивается в 32 раундовых подключа (K1…K8 трижды, затем "
           "K8…K1; при расшифровании — в обратном порядке). 31 раунд с обменом половин и "
           "32-й раунд без обмена.",
 'kuznyechik': "Блочный шифр «Кузнечик» (ГОСТ Р 34.12-2015) на основе SP-сети. Блок 128 бит, "
           "256-битный ключ даёт 10 раундовых ключей. Раунд состоит из наложения ключа "
           "(сложение по модулю 2), нелинейного преобразования S (замена байтов) и "
           "линейного преобразования L. Выполняется 9 раундов и финальное наложение ключа.",
 'rsa':    "Асимметричный шифр RSA. Берутся простые P и Q, вычисляются N = P·Q и функция "
           "Эйлера φ(N) = (P−1)·(Q−1). Открытая экспонента E взаимно проста с φ(N), "
           "секретная D = E⁻¹ mod φ(N). Шифрование: C = Mᴱ mod N; расшифрование: "
           "M = Cᴰ mod N.",
 'elgamal': "Асимметричный шифр Эль-Гамаля, стойкость основана на сложности дискретного "
           "логарифмирования. Открытый ключ y = gˣ mod p. Для каждой буквы выбирается "
           "случайное число k, взаимно простое с φ(p) = p−1: a = g^k mod p, "
           "b = y^k · M mod p. Расшифрование: M = b · (aˣ)⁻¹ mod p.",
 'gost2012': "Цифровая подпись по ГОСТ Р 34.10-2012 на эллиптической кривой над полем "
           "GF(p). Открытый ключ Y = [x]·G (G — базовая точка порядка q). Подпись "
           "сообщения: r = x_P mod q (P = [k]·G), s = (k·H(m) + r·x) mod q. Проверка: "
           "вычисляются z1 = s·h⁻¹, z2 = −r·h⁻¹ mod q, точка P = [z1]·G + [z2]·Y; подпись "
           "верна, если x_P mod q = r.",
 'diffie_hellman': "Протокол выработки общего секретного ключа Диффи — Хеллмана. Открытые ключи "
           "пользователей Y_A = a^{K_A} mod n, Y_B = a^{K_B} mod n. После обмена ими общий "
           "ключ K = Y_B^{K_A} mod n = Y_A^{K_B} mod n = a^{K_A·K_B} mod n. Стойкость "
           "основана на сложности дискретного логарифмирования.",
}

ALGS = [
 ("Блок А. Шифры однозначной замены", "1", "Шифр простой замены АТБАШ", "lab1_atbash.py", "atbash"),
 ("Блок В. Шифры многозначной замены", "7", "S-блок замены ГОСТ Р 34.12-2015 («Магма»)", "lab2_s_block.py", "sblock"),
 ("Блок С. Шифры блочной замены", "8", "Матричный шифр", "lab3_matrix.py", "hill"),
 ("Блок D. Шифры перестановки", "10", "Вертикальная перестановка", "lab4_vertical.py", "vertical"),
 (None, "11", "Решётка Кардано", "lab4_cardano.py", "cardano"),
 (None, "12", "Сеть Фейстеля (ГОСТ 28147-89)", "lab4_feistel.py", "feistel"),
 ("Блок E. Шифры гаммирования", "13", "Одноразовый блокнот К. Шеннона", "lab5_shanon.py", "shannon"),
 (None, "14", "Гаммирование ГОСТ Р 34.13-2015 («Магма», CTR)", "lab5_gamma.py", "gamma"),
 ("Блок F. Поточные шифры", "15", "A5/1", "lab6_a5-1.py", "a51"),
 ("Блок G. Комбинационные шифры", "17", "Магма (ГОСТ Р 34.12-2015)", "lab7_magma.py", "magma"),
 (None, "20", "Кузнечик (ГОСТ Р 34.12-2015)", "lab7_kyznechik.py", "kuznyechik"),
 ("Блок H. Асимметричные шифры", "21", "RSA", "lab8_rsa.py", "rsa"),
 (None, "22", "Эль-Гамаль", "lab8_elgamal.py", "elgamal"),
 ("Блок J. Стандарты цифровых подписей", "27", "ГОСТ Р 34.10-2012", "lab10_gost2012.py", "gost2012"),
 ("Блок K. Обмен ключами", "28", "Обмен ключами по Диффи — Хеллману", "lab11_diffie_hellman.py", "diffie_hellman"),
]

for block, num, name, codefile, key in ALGS:
    if block:
        heading(block, 1)
    heading(f"№ {num}. {name}", 2)
    para("Алгоритм:", bold=True, after=0)
    para(D[key], first_line=1.25)
    heading("Блок-схема программы", 2)
    picture(f"_preview_{key}.png")
    heading("Код программы с комментариями", 2)
    code_block(codefile)
    heading("Тестирование", 2)
    placeholder("Скрин работы программы: шифрование")
    placeholder("Скрин работы программы: расшифрование")
    placeholder("Фото карточки ручного теста с указанием ключа")
    placeholder("Работа с текстом не менее 1000 знаков (шифрование и расшифрование)")
    doc.add_page_break()

# ── ИСПОЛЬЗОВАННЫЕ СТАНДАРТЫ ──────────────────────────────────────────────────
def numbered(items):
    for i, it in enumerate(items, 1):
        p = para(f"{i}. {it}", after=4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)

heading("Использованные стандарты")
numbered([
    "ГОСТ 28147-89. Системы обработки информации. Защита криптографическая. "
    "Алгоритм криптографического преобразования.",
    "ГОСТ Р 34.10-2012. Информационная технология. Криптографическая защита информации. "
    "Процессы формирования и проверки электронной цифровой подписи.",
    "ГОСТ Р 34.11-2012. Информационная технология. Криптографическая защита информации. "
    "Функция хэширования.",
    "ГОСТ Р 34.12-2015. Информационная технология. Криптографическая защита информации. "
    "Блочные шифры («Магма», «Кузнечик»).",
    "ГОСТ Р 34.13-2015. Информационная технология. Криптографическая защита информации. "
    "Режимы работы блочных шифров.",
    "ГОСТ 19.701-90 (ИСО 5807-85). ЕСПД. Схемы алгоритмов, программ, данных и систем. "
    "Условные обозначения и правила выполнения.",
])
doc.add_page_break()

# ── СПИСОК ЛИТЕРАТУРЫ ─────────────────────────────────────────────────────────
heading("Список литературы")
numbered([
    "Программирование криптографических алгоритмов: методические указания к лабораторным "
    "работам (ЛАБКРИПТ_2022). — М.: Московский политехнический университет, 2022.",
    "ГОСТ Р 34.12-2015. Криптографическая защита информации. Блочные шифры.",
    "ГОСТ Р 34.13-2015. Криптографическая защита информации. Режимы работы блочных шифров.",
    "ГОСТ Р 34.10-2012. Криптографическая защита информации. Процессы формирования и проверки "
    "электронной цифровой подписи.",
    "ГОСТ 28147-89. Защита криптографическая. Алгоритм криптографического преобразования.",
    "Шнайер Б. Прикладная криптография: протоколы, алгоритмы и исходные тексты на языке C. — "
    "М.: Триумф, 2002. — 816 с.",
    "Смарт Н. Криптография. — М.: Техносфера, 2005. — 528 с.",
])

_target = "Отчёт_крипта_вар26.docx"
try:
    doc.save(_target)
    print("Готово:", _target)
except PermissionError:
    _alt = "Отчёт_крипта_вар26_NEW.docx"
    doc.save(_alt)
    print(f"ФАЙЛ {_target} ЗАНЯТ (открыт в Word). Сохранил в:", _alt)
